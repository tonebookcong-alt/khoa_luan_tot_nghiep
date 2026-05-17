"""Patch BaoCaoTomTat.docx — apply A1, A3, A5, A6 + rewrite Hạn chế/Hướng phát triển."""
from copy import deepcopy
from docx import Document
from docx.oxml.ns import qn

SRC = 'BaoCaoTomTat.docx'
OUT = 'BaoCaoTomTat.docx'

doc = Document(SRC)


def set_para_text(p, new_text):
    """Replace text in paragraph, preserving format of first run."""
    runs = p.runs
    if not runs:
        p.add_run(new_text)
        return
    runs[0].text = new_text
    # Remove extra runs
    for r in runs[1:]:
        r._element.getparent().remove(r._element)


def append_paragraph_after(anchor_p, text, style=None):
    """Insert a new paragraph immediately after anchor_p, optionally cloning its style."""
    new_p = deepcopy(anchor_p._p)
    # Remove all children except pPr (paragraph properties)
    for child in list(new_p):
        if child.tag != qn('w:pPr'):
            new_p.remove(child)
    anchor_p._p.addnext(new_p)
    # Wrap as Paragraph
    from docx.text.paragraph import Paragraph
    wrap = Paragraph(new_p, anchor_p._parent)
    if style:
        wrap.style = doc.styles[style]
    wrap.add_run(text)
    return wrap


paragraphs = doc.paragraphs

# === A3: Tài nguyên con người + Thời gian ===
set_para_text(paragraphs[105], '- Tài nguyên con người: 4 người')
set_para_text(paragraphs[107], '- Thời gian: Dự án phải được hoàn tất trong vòng 6 tuần (từ 29/03/2026 đến 07/05/2026).')

# === A5: 4.4. Giao diện Admin tổng quan — đổi sang Heading 3 cho khớp TOC ===
paragraphs[178].style = doc.styles['Heading 3']

# === A1: Định hướng giải quyết bài toán ===
set_para_text(
    paragraphs[207],
    '- Phía khách vãng lai: Đăng ký tài khoản, tìm kiếm và lọc tin đăng theo dòng máy/giá/tình trạng, '
    'xem chi tiết sản phẩm kèm biểu đồ giá tham chiếu thị trường và trải nghiệm thử công cụ định giá AI trên ảnh demo.'
)
set_para_text(
    paragraphs[208],
    '- Phía khách hàng (người mua, người bán): Đăng nhập bằng tài khoản nội bộ (JWT) hoặc Google OAuth2; '
    'đăng tin bán máy với mức giá đề xuất tự động từ AI; chat real-time qua Socket.io để thương thảo; '
    'quản lý tin đăng cá nhân, chặn người dùng vi phạm và gửi yêu cầu hỗ trợ đến Admin.'
)
set_para_text(
    paragraphs[209],
    '- Phía admin: Quản lý người dùng (cấp quyền, khóa tài khoản), kiểm duyệt tin đăng, quản lý danh mục điện thoại '
    'và bảng hệ số khấu hao; theo dõi KPI tổng quan (số người dùng, số tin đăng, biến động giá theo dòng máy) '
    'qua biểu đồ Recharts; phản hồi yêu cầu hỗ trợ từ người dùng.'
)

# === Hạn chế: rewrite + thêm 3 mục mới ===
set_para_text(
    paragraphs[242],
    '- Phạm vi dataset thị giác: Mô hình YOLOv11s hiện chỉ nhận diện 9 thế hệ iPhone '
    '(từ iPhone 6 đến iPhone 17), chưa hỗ trợ các dòng Android (Samsung, Xiaomi, OPPO…) do giới hạn '
    'thời gian thu thập, label dữ liệu và chi phí GPU.'
)
set_para_text(
    paragraphs[243],
    '- Phạm vi định giá: Công thức P_final chỉ áp dụng từ iPhone X trở lên; các dòng cũ hơn '
    '(iPhone 6, 7, 8) chỉ trả giá tham khảo do thị trường nhiều tin rác và scam, khiến median P_market không đáng tin cậy.'
)

# Insert thêm 3 paragraph hạn chế sau paragraph 243
anchor = paragraphs[243]
p = append_paragraph_after(
    anchor,
    '- Đánh giá tình trạng pin: Không thể nhận diện chai pin từ ảnh, hệ thống phụ thuộc vào input '
    'người bán khai báo nên có khả năng bị sai lệch nếu người bán không trung thực.',
    style='List Paragraph',
)
p = append_paragraph_after(
    p,
    '- Thu thập giá thị trường: AI Agent market-scraping ở giai đoạn báo cáo dùng dữ liệu mẫu có cấu trúc; '
    'chưa chạy crawler real-time liên tục do giới hạn quota của Apify free tier và rủi ro bị Chợ Tốt chặn IP.',
    style='List Paragraph',
)
p = append_paragraph_after(
    p,
    '- Cơ chế giao dịch: Hệ thống mới dừng lại ở việc kết nối hai bên qua Chat Real-time, '
    'chưa tích hợp thanh toán đảm bảo (escrow) nên người dùng vẫn phải tự chốt giao dịch off-platform.',
    style='List Paragraph',
)

# Cập nhật danh sách paragraphs sau khi insert (paragraphs cũ vẫn còn ref, chỉ cần re-fetch nếu cần index)
paragraphs = doc.paragraphs

# === Hướng phát triển: rewrite 3 mục cũ + thêm 3 mục mới ===
# Find lại index bằng text (vì index đã shift)
idx_huong = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == 'Hướng phát triển:':
        idx_huong = i
        break

# 3 bullet ngay sau "Hướng phát triển:" → rewrite
set_para_text(
    doc.paragraphs[idx_huong + 1],
    '- Mở rộng dataset thị giác sang điện thoại Android (Samsung Galaxy S/Note, Xiaomi, OPPO, vivo) '
    'để phục vụ phân khúc thị trường rộng hơn ngoài iPhone.'
)
set_para_text(
    doc.paragraphs[idx_huong + 2],
    '- Tích hợp cơ chế thanh toán đảm bảo (escrow) qua VNPay/Momo/ZaloPay để tạm giữ tiền 7 ngày, '
    'bảo vệ quyền lợi cả hai bên trong giao dịch máy cũ.'
)
set_para_text(
    doc.paragraphs[idx_huong + 3],
    '- Đưa AI Agent market-scraping vào vận hành thực tế với crawler theo lịch (cron) trên VPS, '
    'kết hợp proxy rotation để cập nhật P_market hàng ngày từ Chợ Tốt và các marketplace khác.'
)

# Thêm 3 mục mới sau idx_huong + 3
anchor = doc.paragraphs[idx_huong + 3]
p = append_paragraph_after(
    anchor,
    '- Tích hợp API kiểm tra pin/IMEI qua dịch vụ bên thứ ba (iCheck, CheckPhone…) để cross-check thông tin '
    'người bán khai báo, giảm phụ thuộc vào input thủ công.',
    style='List Paragraph',
)
p = append_paragraph_after(
    p,
    '- Phát triển ứng dụng mobile (React Native) tận dụng camera native để chụp ảnh đa góc, '
    'cải thiện chất lượng input cho mô hình thị giác và nâng cao độ chính xác định giá.',
    style='List Paragraph',
)
p = append_paragraph_after(
    p,
    '- Mở rộng pricing engine cho các dòng iPhone đời cũ (iPhone 6, 7, 8) khi đã thu thập đủ dữ liệu '
    'thị trường đáng tin cậy và lọc được phần lớn tin rác.',
    style='List Paragraph',
)

# === A6: Tài liệu tham khảo ===
# Find TÀI LIỆU THAM KHẢO heading
idx_tltk = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == 'TÀI LIỆU THAM KHẢO':
        idx_tltk = i
        break

refs = [
    '[1] NestJS Documentation. https://docs.nestjs.com/ (truy cập tháng 05/2026).',
    '[2] Next.js Documentation. https://nextjs.org/docs (truy cập tháng 05/2026).',
    '[3] Prisma ORM Documentation. https://www.prisma.io/docs (truy cập tháng 05/2026).',
    '[4] Mongoose ODM Documentation. https://mongoosejs.com/docs/ (truy cập tháng 05/2026).',
    '[5] Redis Documentation. https://redis.io/docs/ (truy cập tháng 05/2026).',
    '[6] Ultralytics. YOLOv11 Documentation. https://docs.ultralytics.com/models/yolo11/ (truy cập tháng 05/2026).',
    '[7] LangChain.js Documentation. https://js.langchain.com/docs/ (truy cập tháng 05/2026).',
    '[8] Socket.IO Documentation. https://socket.io/docs/v4/ (truy cập tháng 05/2026).',
    '[9] Google Developers. Using OAuth 2.0 to Access Google APIs. https://developers.google.com/identity/protocols/oauth2 (truy cập tháng 05/2026).',
    '[10] Roboflow. Computer Vision Annotation & Training Platform. https://roboflow.com/ (truy cập tháng 05/2026).',
    '[11] Google AI. Gemini API Documentation. https://ai.google.dev/gemini-api/docs (truy cập tháng 05/2026).',
    '[12] Docker Documentation. https://docs.docker.com/ (truy cập tháng 05/2026).',
]

anchor = doc.paragraphs[idx_tltk]
for ref in refs:
    anchor = append_paragraph_after(anchor, ref, style='Normal')

doc.save(OUT)
print(f'OK — saved {OUT}')
print(f'Total paragraphs now: {len(doc.paragraphs)}')
